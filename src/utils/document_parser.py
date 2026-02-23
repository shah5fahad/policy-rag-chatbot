from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
import asyncio, re, os
from pypdf import PdfReader  # Use pypdf instead of PyPDF2
from concurrent.futures import ThreadPoolExecutor
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.opc.exceptions import PackageNotFoundError
from dotenv import load_dotenv
from loguru import logger
from PIL import Image
from pathlib import Path
import pytesseract

load_dotenv()
CHUNK_SIZE = int(os.getenv("DOCUMENT_CHUNK_SIZE", 1000))  # Number of words per chunk for text splitting
CHUNK_OVERLAP = 100     # overlap for RAG quality


class DocumentParser(ABC):
    """Base abstract class for document parsers."""
    
    @abstractmethod
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse document and return extracted content."""
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        pass


class TextParser(DocumentParser):
    """Parser for plain text files."""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            content = await loop.run_in_executor(
                executor,
                self._parse_sync,
                file_path
            )
        return content
    
    def _parse_sync(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {
            "text": text,
            "chunks": self._chunk_text(text),
            "metadata": {"char_count": len(text), "word_count": len(text.split())}
        }
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks for embedding."""
        words = text.split()
        chunks = []
        for i in range(0, len(words), CHUNK_SIZE):
            chunk = ' '.join(words[i:i+CHUNK_SIZE])
            chunks.append(chunk)
        return chunks
    
    def get_supported_extensions(self) -> List[str]:
        return ['txt', 'md']


class PDFParser:
    """Parser for PDF files with robust error handling."""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF file asynchronously."""
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            content = await loop.run_in_executor(
                executor,
                self._parse_sync,
                file_path
            )
        return content
    
    def _parse_sync(self, file_path: str) -> Dict[str, Any]:
        """Synchronous PDF parsing with comprehensive error handling."""
        try:
            # Validate file
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            if Path(file_path).suffix.lower() != '.pdf':
                raise ValueError(f"Not a PDF file: {file_path}")
            
            text = ""
            metadata = {}
            page_count = 0
            extraction_errors = []
            
            try:
                # Use PdfReader from pypdf
                with open(file_path, 'rb') as f:
                    pdf_reader = PdfReader(f)
                    page_count = len(pdf_reader.pages)
                    
                    # Extract metadata safely
                    try:
                        if pdf_reader.metadata:
                            metadata = {}
                            for key, value in pdf_reader.metadata.items():
                                if value:
                                    clean_key = key.replace('/', '').lower()
                                    metadata[clean_key] = str(value)
                    except Exception as e:
                        metadata = {"note": f"Metadata extraction failed: {str(e)}"}
                    
                    # Extract text from each page
                    for page_num in range(page_count):
                        try:
                            page = pdf_reader.pages[page_num]
                            
                            # Try different extraction methods
                            page_text = ""
                            
                            # Method 1: Standard extraction
                            try:
                                page_text = page.extract_text() or ""
                            except:
                                pass
                            
                            # Method 2: If no text, try extracting with different settings
                            if not page_text.strip():
                                try:
                                    page_text = page.extract_text(extraction_mode="layout") or ""
                                except:
                                    pass
                            
                            # Clean text if we got any
                            if page_text.strip():
                                # Clean the text
                                page_text = re.sub(r'\s+', ' ', page_text)
                                page_text = re.sub(r'(\w)-\n(\w)', r'\1\2', page_text)
                                page_text = re.sub(r'[ \t]+', ' ', page_text)
                                page_text = re.sub(r'\n\s+\n', '\n\n', page_text)
                                
                                text += f"\n\n===== Page {page_num + 1} =====\n\n{page_text}\n"
                            else:
                                # For scanned/image-based PDFs or tables
                                text += f"\n\n===== Page {page_num + 1} =====\n[Content may be image-based or table format - text extraction limited]\n"
                                
                        except Exception as e:
                            error_msg = f"Page {page_num + 1} failed: {str(e)}"
                            extraction_errors.append(error_msg)
                            text += f"\n\n===== Page {page_num + 1} =====\n[Error: Could not extract text - {str(e)}]\n"
                            
            except Exception as e:
                # If PDF is completely unreadable
                return {
                    "text": f"Error reading PDF: {str(e)}",
                    "chunks": [],
                    "metadata": {
                        "filename": os.path.basename(file_path),
                        "error": str(e),
                        "page_count": 0
                    }
                }
            
            # Clean up the full text
            text = re.sub(r'\n{4,}', '\n\n', text)
            
            # Build result
            result = {
                "text": text,
                "chunks": self._chunk_text(text),
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "page_count": page_count,
                    "pdf_metadata": metadata,
                    "extraction_errors": extraction_errors if extraction_errors else None,
                    "has_content": bool(text.strip() and "[Content may be image-based" not in text)
                }
            }
            
            logger.info(f"Parsed {file_path} with {page_count} pages")
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {str(e)}")
            return {
                "text": "",
                "chunks": [],
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "error": str(e),
                    "page_count": 0
                }
            }
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if not text or not text.strip():
            return []
        
        # Remove the page markers for chunking to get cleaner chunks
        clean_text = re.sub(r'===== Page \d+ =====\n\n', '', text)
        
        words = clean_text.split()
        chunks = []
        
        start = 0
        while start < len(words):
            end = min(start + CHUNK_SIZE, len(words))
            chunk = ' '.join(words[start:end])
            if chunk.strip():
                chunks.append(chunk)
            
            if end == len(words):
                break
            start += CHUNK_SIZE - CHUNK_OVERLAP
        
        return chunks
    
    def get_supported_extensions(self) -> List[str]:
        return ['pdf']


class DocxParser(DocumentParser):
    """
    Industrial-grade DOCX parser supporting:
    - paragraphs
    - tables
    - headers / footers
    - metadata
    - robust chunking
    """

    async def parse(self, file_path: str) -> Dict[str, Any]:
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            content = await loop.run_in_executor(
                executor,
                self._parse_sync,
                file_path
            )
        return content
    
    def _parse_sync(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = DocxDocument(file_path)
        except PackageNotFoundError:
            raise ValueError("Invalid or corrupted DOCX file")

        elements = []

        # Extract main body content
        for block in self._iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    elements.append(text)

            elif isinstance(block, Table):
                table_text = self._extract_table(block)
                if table_text:
                    elements.append(table_text)

        # Extract headers & footers
        for section in doc.sections:
            header = section.header
            footer = section.footer

            for para in header.paragraphs:
                if para.text.strip():
                    elements.append(para.text.strip())

            for para in footer.paragraphs:
                if para.text.strip():
                    elements.append(para.text.strip())

        full_text = "\n\n".join(elements)

        metadata = self._extract_metadata(doc, file_path, full_text)

        return {
            "text": full_text,
            "chunks": self._chunk_text(full_text),
            "metadata": metadata,
        }

    def _iter_block_items(self, parent):
        """
        Yield paragraphs and tables in document order.
        """
        for child in parent.element.body.iterchildren():
            if child.tag.endswith('}p'):
                yield Paragraph(child, parent)
            elif child.tag.endswith('}tbl'):
                yield Table(child, parent)

    def _extract_table(self, table: Table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _extract_metadata(self, doc, file_path: str, text: str) -> Dict[str, Any]:
        core = doc.core_properties

        return {
            "filename": os.path.basename(file_path),
            "title": core.title,
            "author": core.author,
            "subject": core.subject,
            "created": str(core.created) if core.created else None,
            "modified": str(core.modified) if core.modified else None,
            "paragraph_count": len(doc.paragraphs),
            "word_count": len(text.split()),
            "character_count": len(text),
        }
        
    def _chunk_text(self, text: str) -> List[str]:
        """
        Smart overlapping chunking for RAG systems.
        """
        words = text.split()
        chunks = []

        start = 0
        while start < len(words):
            end = start + CHUNK_SIZE
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
            start += CHUNK_SIZE - CHUNK_OVERLAP

        return chunks
    
    def get_supported_extensions(self) -> List[str]:
        return ['docx']


class ImageParser(DocumentParser):
    """Parser for images using OCR."""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            content = await loop.run_in_executor(
                executor,
                self._parse_sync,
                file_path
            )
        return content
    
    def _parse_sync(self, file_path: str) -> Dict[str, Any]:
        # Open image
        image = Image.open(file_path)
        
        # Perform OCR
        text = pytesseract.image_to_string(image)
        
        return {
            "text": text,
            "chunks": self._chunk_text(text),
            "metadata": {
                "image_size": image.size,
                "image_format": image.format,
                "image_mode": image.mode
            }
        }
    
    def _chunk_text(self, text: str) -> List[str]:
        sentences = re.split(r'(?<=[.!?]) +', text)
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            if current_size + sentence_words > CHUNK_SIZE and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = sentence_words
            else:
                current_chunk.append(sentence)
                current_size += sentence_words
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def get_supported_extensions(self) -> List[str]:
        return ['jpg', 'jpeg', 'png']


class ParserFactory:
    """Factory to get appropriate parser for file type."""
    
    def __init__(self):
        self.parsers = [
            PDFParser(),
            DocxParser(),
            TextParser(),
            ImageParser()
        ]
        self.parser_map = {}
        for parser in self.parsers:
            for ext in parser.get_supported_extensions():
                self.parser_map[ext] = parser
    
    def get_parser(self, file_extension: str) -> Optional[DocumentParser]:
        """Get parser for file extension."""
        return self.parser_map.get(file_extension.lower())